#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""每日复盘 Word 报告生成器。

报告章节：
1. 封面（日期、总览数据）
2. 每日情绪指数（溢价率/炸板率/晋级率/生命周期）
3. 涨停数据（全部涨停股：代码/名称/板数/涨停原因/所属概念）
4. 龙虎榜数据（全部龙虎榜个股详情）
5. 炸板分析（曾涨停个股）
6. 概念分析（TOP概念+成分股明细）
7. 当前最高板
8. 市场梯队诊断（以每个板级为锚点，展示全部板块梯队）
"""
import os
import struct
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = Path(__file__).parent.resolve()
REPORT_DIR = BASE_DIR / 'reports'
os.makedirs(REPORT_DIR, exist_ok=True)


# ═══════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════

def _fmt_amt(val):
    """格式化金额（万元→可读）"""
    if abs(val) >= 10000:
        return '%.1f亿' % (val / 10000)
    return '%d万' % int(val)


def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" w:color="{val.get("color","D9D9D9")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_styled_table(doc, headers, rows, col_widths=None,
                      header_color='1A1A2E', row_text_color=None):
    """创建统一风格的表格。row_text_color: (R,G,B) 元组，覆盖全部数据行字体颜色。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        _set_cell_shading(cell, header_color)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c < 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if row_text_color:
                run.font.color.rgb = RGBColor(*row_text_color)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            # 交替行背景
            if r % 2 == 1:
                _set_cell_shading(cell, 'F6F8FA')

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


def _add_section_title(doc, title, number=None):
    """添加带编号的章节标题"""
    if number:
        full = '%s  %s' % (number, title)
    else:
        full = title
    p = doc.add_paragraph()
    run = p.add_run(full)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    # 下划线装饰
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="E94560"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_sub_title(doc, title, color=None):
    """添加子标题"""
    if color is None:
        color = (0xE9, 0x45, 0x60)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def _add_stat_box(doc, items):
    """添加统计卡片行"""
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, color) in enumerate(items):
        cell = table.rows[0].cells[i]
        cell.text = ''
        # 数值
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(value))
        r1.font.size = Pt(18)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(*_hex_to_rgb(color))
        r1.font.name = 'Microsoft YaHei'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        # 标签
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r2.font.name = 'Microsoft YaHei'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        _set_cell_shading(cell, '161B22')
    doc.add_paragraph()


def _hex_to_rgb(hex_str):
    """#RRGGBB → (R, G, B)"""
    h = hex_str.lstrip('#')
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def _find_reason(code, up_stocks):
    """从涨停股列表中查找某只股票的涨停原因。

    Args:
        code: 6位股票代码。
        up_stocks: 涨停股列表 [{code, name, reason_type, ...}]。

    Returns:
        涨停原因字符串，找不到返回空字符串。
    """
    for s in up_stocks:
        if str(s.get('code', '')).zfill(6) == code:
            return s.get('reason_type', '') or s.get('reason', '') or ''
    return ''


# ═══════════════════════════════════════════════════════
# 报告生成主函数
# ═══════════════════════════════════════════════════════

def generate_word_report(
    date_str,
    up_stocks,
    board_days,
    em_data,
    lhb_detail,
    concept_result,
    premium_rate,
    zhaban_rate,
    code_name_map,
    callback=None,
    yesterday_zt_count=0,
):
    """生成每日复盘 Word 报告。

    Args:
        date_str: 日期字符串 YYYYMMDD
        up_stocks: 涨停股列表 [{code, name, reason_type, ...}]
        board_days: {code: 连板天数}
        em_data: 东财数据 {zting, zhaban, all_dting, cengdting}
        lhb_detail: 龙虎榜详情 {code: {name, net_buy, buy_total, sell_total, brokers}}
        concept_result: 概念分析结果 {ranked, top3, code_concepts}
        premium_rate: 溢价率(%)
        zhaban_rate: 炸板率(%)
        code_name_map: {code: name} 全量名称映射
        callback: 日志回调

    Returns:
        生成的 docx 文件路径
    """
    log = callback or print

    # ── 数据预处理 ──
    d_display = '%s-%s-%s' % (date_str[:4], date_str[4:6], date_str[6:8])

    # 涨停分类
    zt_by_board = {}  # {board: [code, ...]}
    sb_codes = []
    lb_codes = []
    for s in up_stocks:
        code = str(s.get('code', '')).zfill(6)
        bd = board_days.get(code, 1)
        zt_by_board.setdefault(bd, []).append(code)
        if bd == 1:
            sb_codes.append(code)
        else:
            lb_codes.append(code)

    max_board = max(board_days.values()) if board_days else 0
    sorted_boards = sorted(zt_by_board.keys(), reverse=True)

    # 炸板
    zhaban_codes = em_data.get('zhaban', [])

    # 概念排名
    ranked = concept_result.get('ranked', [])
    code_concepts = concept_result.get('code_concepts', {})

    # 晋级率 = 今日连板数 / 昨日涨停数
    # 若未传入昨日涨停数则降级为连板占比（连板数/今日涨停数）
    promo_rate = 0
    if lb_codes:
        if yesterday_zt_count > 0:
            promo_rate = len(lb_codes) / yesterday_zt_count * 100
        else:
            promo_rate = len(lb_codes) / len(up_stocks) * 100 if up_stocks else 0

    # 生命周期
    zt_count = len(up_stocks)
    dt_count = len(em_data.get('all_dting', []))
    lhb_count = len(lhb_detail)
    ck = sum([zt_count >= 5, max_board >= 4, len(lb_codes) >= 3])
    if ck >= 3:
        life_cycle = '爆发期'
        life_detail = '主升阶段，梯队完整'
    elif ck >= 2:
        life_cycle = '发酵期'
        life_detail = '资金开始关注'
    else:
        life_cycle = '分歧/退潮期'
        life_detail = '热度下降'

    log('  生成Word报告: %s (%d涨停/%d跌停/%d龙虎榜)' % (d_display, zt_count, dt_count, lhb_count))

    # ── 创建文档 ──
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ════════════════════════════════════════
    # 1. 封面
    # ════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('每日复盘报告')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('V3.1 · 通达信自动复盘系统')
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    r2.font.name = 'Microsoft YaHei'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(d_display)
    r3.font.size = Pt(20)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r3.font.name = 'Microsoft YaHei'
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # 封面摘要
    cover_items = [
        ('涨停', str(zt_count), '#F85149'),
        ('跌停', str(dt_count), '#58A6FF'),
        ('最高板', '%d板' % max_board, '#E3B341'),
        ('龙虎榜', str(lhb_count), '#3FB950'),
        ('生命周期', life_cycle, '#E94560'),
    ]
    _add_stat_box(doc, cover_items)

    p_time = doc.add_paragraph()
    p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_time = p_time.add_run('生成时间: %s' % datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    r_time.font.size = Pt(9)
    r_time.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    r_time.font.name = 'Microsoft YaHei'
    r_time._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. 每日情绪指数
    # ════════════════════════════════════════
    _add_section_title(doc, '每日情绪指数', '一')
    _add_stat_box(doc, [
        ('溢价率 (昨涨停今开盘)', '%.1f%%' % premium_rate, '#FFD93D'),
        ('炸板率', '%.0f%%' % zhaban_rate, '#6BCB77'),
        ('晋级率 (首板→连板)', '%.0f%%' % promo_rate, '#FF6B6B'),
        ('生命周期', '%s — %s' % (life_cycle, life_detail), '#E3B341'),
    ])

    # 情绪解读
    _add_sub_title(doc, '情绪解读')
    interpretations = []
    if premium_rate > 3:
        interpretations.append('溢价率%.1f%%，昨日涨停股今日高开明显，市场追涨意愿强。' % premium_rate)
    elif premium_rate > 0:
        interpretations.append('溢价率%.1f%%，昨日涨停股小幅高开，情绪中性偏暖。' % premium_rate)
    else:
        interpretations.append('溢价率%.1f%%，昨日涨停股低开，市场追涨意愿弱。' % premium_rate)

    if zhaban_rate > 40:
        interpretations.append('炸板率%.0f%%，封板成功率低，注意追高风险。' % zhaban_rate)
    elif zhaban_rate > 20:
        interpretations.append('炸板率%.0f%%，封板成功率一般。' % zhaban_rate)
    else:
        interpretations.append('炸板率%.0f%%，封板成功率高，资金封板意愿强。' % zhaban_rate)

    if life_cycle == '爆发期':
        interpretations.append('市场处于爆发期，梯队完整，可积极关注主线龙头。')
    elif life_cycle == '发酵期':
        interpretations.append('市场处于发酵期，关注是否能进一步发酵为爆发期。')
    else:
        interpretations.append('市场处于分歧/退潮期，建议控制仓位，等待明确信号。')

    for text in interpretations:
        p = doc.add_paragraph()
        r = p.add_run('• ' + text)
        r.font.size = Pt(10)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 3. 涨停数据
    # ════════════════════════════════════════
    _add_section_title(doc, '涨停数据（%d只）' % zt_count, '二')

    # 按板数分组展示
    for board in sorted_boards:
        codes = zt_by_board.get(board, [])
        if not codes:
            continue
        if board == 1:
            label = '首板（%d只）' % len(codes)
        elif board == 2:
            label = '2板（%d只）' % len(codes)
        elif board == 3:
            label = '3板（%d只）' % len(codes)
        elif board == 4:
            label = '4板（%d只）' % len(codes)
        else:
            label = '%d板（%d只）' % (board, len(codes))

        _add_sub_title(doc, label)

        rows = []
        for code in codes:
            name = code_name_map.get(code, '')
            reason = _find_reason(code, up_stocks)
            concepts_str = ''
            concepts = code_concepts.get(code, [])
            concepts_str = '、'.join(concepts[:3]) if concepts else '-'
            if len(reason) > 30:
                reason = reason[:28] + '..'
            rows.append([code, name, reason, concepts_str])

        _add_styled_table(doc,
                          ['代码', '名称', '涨停原因', '所属概念'],
                          rows,
                          col_widths=[2.0, 2.0, 6.0, 6.0])

    doc.add_page_break()

    # ════════════════════════════════════════
    # 4. 龙虎榜数据（全部）
    # ════════════════════════════════════════
    _add_section_title(doc, '龙虎榜数据（%d只）' % lhb_count, '三')

    if lhb_detail:
        # 按净买额排序
        sorted_lhb = sorted(lhb_detail.items(),
                            key=lambda x: float(x[1].get('net_buy', 0)), reverse=True)

        for code, item in sorted_lhb:
            name = item.get('name', '') or code_name_map.get(code, '')
            net_buy = item.get('net_buy', 0)
            buy_total = item.get('buy_total', 0)
            sell_total = item.get('sell_total', 0)
            brokers = item.get('brokers', {})

            # 子标题
            sign = '+' if net_buy >= 0 else ''
            title = '%s  %s  净买：%s%s  买入：%s  卖出：%s' % (
                code, name,
                sign, _fmt_amt(net_buy),
                _fmt_amt(buy_total), _fmt_amt(sell_total)
            )
            _add_sub_title(doc, title)

            # 买入席位
            buy_list = brokers.get('buy', [])[:5]
            sell_list = sorted(brokers.get('sell', []),
                               key=lambda x: float(x.get('sell', 0)), reverse=True)[:5]

            # 合并席位表格
            max_len = max(len(buy_list), len(sell_list))
            rows = []
            for i in range(max_len):
                buy_name = ''
                buy_amt = ''
                sell_name = ''
                sell_amt = ''
                if i < len(buy_list):
                    b = buy_list[i]
                    # 缩写券商名（防None）
                    buy_name_raw = str(b.get('name') or '')
                    for s_str in ['证券', '有限', '责任', '股份', '公司', '分公司', '营业部']:
                        buy_name_raw = buy_name_raw.replace(s_str, '')
                    buy_name = buy_name_raw[:8]
                    buy_amt = _fmt_amt(float(b.get('buy') or 0))
                if i < len(sell_list):
                    b = sell_list[i]
                    sell_name_raw = str(b.get('name') or '')
                    for s_str in ['证券', '有限', '责任', '股份', '公司', '分公司', '营业部']:
                        sell_name_raw = sell_name_raw.replace(s_str, '')
                    sell_name = sell_name_raw[:8]
                    sell_amt = _fmt_amt(float(b.get('sell') or 0))
                rows.append([buy_name, buy_amt, sell_name, sell_amt])

            _add_styled_table(doc,
                              ['买入席位', '买入金额', '卖出席位', '卖出金额'],
                              rows,
                              col_widths=[4.5, 2.5, 4.5, 2.5],
                              header_color='E94560')
    else:
        p = doc.add_paragraph()
        r = p.add_run('当日无龙虎榜数据')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 5. 炸板分析
    # ════════════════════════════════════════
    zhaban_count = len(zhaban_codes)
    _add_section_title(doc, '炸板分析（%d只，炸板率%.0f%%）' % (zhaban_count, zhaban_rate), '四')

    if zhaban_codes:
        rows = []
        for code in zhaban_codes:
            name = code_name_map.get(str(code).zfill(6), '')
            rows.append([str(code).zfill(6), name])
        _add_styled_table(doc,
                          ['代码', '名称'],
                          rows,
                          col_widths=[3.0, 4.0])
    else:
        p = doc.add_paragraph()
        r = p.add_run('当日无炸板')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 6. 概念分析
    # ════════════════════════════════════════
    _add_section_title(doc, '概念分析', '五')

    if ranked:
        for i, item in enumerate(ranked[:10]):
            cnt = item[0]
            concept = item[1]
            codes = item[2] if len(item) >= 3 else []

            rank_icon = ['🥇', '🥈', '🥉'][i] if i < 3 else '%d.' % (i + 1)
            _add_sub_title(doc, '%s  %s（%d只涨停）' % (rank_icon, concept, cnt))

            rows = []
            for code in codes:
                name = code_name_map.get(code, '')
                bd = board_days.get(code, 1)
                rows.append([code, name, '%d板' % bd])

            _add_styled_table(doc,
                              ['代码', '名称', '板数'],
                              rows,
                              col_widths=[3.0, 3.0, 2.0])

            # 概念间加小间距
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph()
        r = p.add_run('当日无概念数据')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 7. 当前最高板
    # ════════════════════════════════════════
    _add_section_title(doc, '当前最高板 — %d板' % max_board, '六')

    if max_board > 0:
        top_codes = zt_by_board.get(max_board, [])
        for code in top_codes:
            name = code_name_map.get(code, '')
            reason = _find_reason(code, up_stocks)
            concepts = code_concepts.get(code, [])
            concepts_str = '、'.join(concepts[:5]) if concepts else '-'

            p = doc.add_paragraph()
            r = p.add_run('● %s  %s  %d板' % (code, name, max_board))
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xF8, 0x51, 0x49)
            r.font.name = 'Microsoft YaHei'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            p2 = doc.add_paragraph()
            r2 = p2.add_run('  涨停原因: %s' % (reason or '-'))
            r2.font.size = Pt(10)
            r2.font.name = 'Microsoft YaHei'
            r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            p3 = doc.add_paragraph()
            r3 = p3.add_run('  所属概念: %s' % concepts_str)
            r3.font.size = Pt(10)
            r3.font.name = 'Microsoft YaHei'
            r3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

            # 是否在龙虎榜
            if code in lhb_detail:
                lhb_item = lhb_detail[code]
                net_buy = lhb_item.get('net_buy', 0)
                p4 = doc.add_paragraph()
                r4 = p4.add_run('  龙虎榜: 净买%s' % _fmt_amt(net_buy))
                r4.font.size = Pt(10)
                r4.font.color.rgb = RGBColor(0x3F, 0xB9, 0x50)
                r4.font.name = 'Microsoft YaHei'
                r4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        p = doc.add_paragraph()
        r = p.add_run('当日无涨停数据')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 8. 市场梯队诊断
    # ════════════════════════════════════════
    _add_section_title(doc, '市场梯队诊断', '七')

    # 按板级从高到低诊断
    for board in sorted_boards:
        if board < 2:
            break  # 首板另做
        codes = zt_by_board.get(board, [])
        if not codes:
            continue

        label = '%d板梯队（%d只）' % (board, len(codes)) if board > 1 else '首板（%d只）' % len(codes)
        _add_sub_title(doc, label)

        rows = []
        for code in codes:
            name = code_name_map.get(code, '')
            reason = _find_reason(code, up_stocks)
            if len(reason) > 20:
                reason = reason[:18] + '..'
            concepts = code_concepts.get(code, [])
            concepts_str = '、'.join(concepts[:4]) if concepts else '-'

            # 龙虎榜标记
            lhb_mark = ''
            if code in lhb_detail:
                net_buy = lhb_detail[code].get('net_buy', 0)
                lhb_mark = '龙虎净买%s' % _fmt_amt(net_buy)

            rows.append([code, name, reason, concepts_str, lhb_mark])

        _add_styled_table(doc,
                          ['代码', '名称', '涨停原因', '所属概念', '龙虎榜'],
                          rows,
                          col_widths=[2.0, 2.0, 3.5, 4.5, 3.5])

    # 首板梯队
    sb_codes_list = zt_by_board.get(1, [])
    if sb_codes_list:
        _add_sub_title(doc, '首板梯队（%d只）' % len(sb_codes_list))

        rows = []
        for code in sb_codes_list:
            name = code_name_map.get(code, '')
            reason = _find_reason(code, up_stocks)
            if len(reason) > 20:
                reason = reason[:18] + '..'
            concepts = code_concepts.get(code, [])
            concepts_str = '、'.join(concepts[:4]) if concepts else '-'
            lhb_mark = ''
            if code in lhb_detail:
                net_buy = lhb_detail[code].get('net_buy', 0)
                lhb_mark = '龙虎净买%s' % _fmt_amt(net_buy)
            rows.append([code, name, reason, concepts_str, lhb_mark])

        _add_styled_table(doc,
                          ['代码', '名称', '涨停原因', '所属概念', '龙虎榜'],
                          rows,
                          col_widths=[2.0, 2.0, 3.5, 4.5, 3.5])

    # ── 免责声明 ──
    doc.add_page_break()
    _add_section_title(doc, '免责声明', '八')
    disclaimer = (
        '本报告由复盘助手 V3.1 自动生成，数据来源于同花顺、东方财富、ClawHub 等公开数据平台。\n\n'
        '⚠️ 免责声明：本软件为数据整理工具，仅提供公开市场数据的自动分类与展示功能。'
        '不提供任何投资建议、不推荐任何股票、不预测市场走势。'
        '晋级率、溢价率、炸板率等指标仅为历史数据统计，不构成投资建议。'
        '用户基于本报告做出的任何投资决策，均由用户自行承担风险。\n\n'
        '股市有风险，投资需谨慎。'
    )
    p = doc.add_paragraph()
    r = p.add_run(disclaimer)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ── 保存 ──
    filename = '复盘报告_%s.docx' % d_display
    filepath = REPORT_DIR / filename
    doc.save(str(filepath))
    log('  Word报告已保存: %s' % filepath)

    return str(filepath)


# ═══════════════════════════════════════════════════════
# 命令行入口
# ═══════════════════════════════════════════════════════

if __name__ == '__main__':
    # 测试：从 latest.json 读取数据生成报告
    import json
    from utils import load_json, DATA_DIR

    latest = load_json(DATA_DIR / 'latest.json')
    if not latest:
        print('latest.json 不存在，请先运行 app.py --update')
        exit(1)

    date_str = latest.get('date', datetime.now().strftime('%Y%m%d'))
    print('从 latest.json 生成报告: %s' % date_str)

    # 加载历史历史数据获取更多详情
    history_file = DATA_DIR / 'history' / ('%s.json' % date_str)
    history = load_json(history_file) if history_file.exists() else {}

    # 简化测试（仅用latest数据）
    up_stocks = []
    board_days = {}
    em_data = {'zting': [], 'zhaban': [], 'all_dting': [], 'cengdting': []}
    lhb_detail = {}
    concept_result = {'ranked': [], 'top3': [], 'code_concepts': {}}
    code_name_map = {}

    path = generate_word_report(
        date_str, up_stocks, board_days, em_data, lhb_detail,
        concept_result, latest.get('premiumRate', 0),
        latest.get('zhabanRate', 0), code_name_map
    )
    print('报告: %s' % path)
